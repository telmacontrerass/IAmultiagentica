"""Unit tests for ci2lab/harness/tools/docx_writer.py."""

from __future__ import annotations

from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def tmp_workspace(tmp_path: Path) -> Path:
    """Temporary workspace with a minimal template."""
    return tmp_path


@pytest.fixture()
def simple_template(tmp_workspace: Path) -> Path:
    """Create a template DOCX with basic placeholders."""
    pytest.importorskip("docx", reason="python-docx not installed")
    from docx import Document

    tpl_dir = tmp_workspace / "templates"
    tpl_dir.mkdir()
    tpl = tpl_dir / "template.docx"

    doc = Document()
    doc.add_paragraph("Name: {{name}}")
    doc.add_paragraph("Date: {{date}}")
    doc.add_paragraph("No placeholder")
    # Table with placeholder
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("{{city}}")
    table.rows[0].cells[1].paragraphs[0].add_run("description")
    doc.save(str(tpl))
    return tpl


# ---------------------------------------------------------------------------
# Tests de fill_docx_template
# ---------------------------------------------------------------------------


class TestFillDocxTemplate:
    def test_basic_substitution(self, tmp_workspace: Path, simple_template: Path) -> None:
        """Substitute placeholders in paragraphs and tables."""
        from docx import Document

        from ci2lab.harness.tools.docx_writer import fill_docx_template

        output = tmp_workspace / "output.docx"
        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="output.docx",
            fields={"name": "Clara", "date": "2026-06-14", "city": "Madrid"},
        )

        assert "Error" not in result
        assert output.is_file()

        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs]
        assert "Name: Clara" in texts
        assert "Date: 2026-06-14" in texts
        assert "No placeholder" in texts

        # Table
        cell_text = doc.tables[0].rows[0].cells[0].paragraphs[0].text
        assert cell_text == "Madrid"

    def test_no_fields(self, tmp_workspace: Path, simple_template: Path) -> None:
        """With no fields the template is copied unchanged."""
        from docx import Document

        from ci2lab.harness.tools.docx_writer import fill_docx_template

        output = tmp_workspace / "no_fields.docx"
        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="no_fields.docx",
            fields={},
        )

        assert "Error" not in result
        doc = Document(str(output))
        assert any("{{name}}" in p.text for p in doc.paragraphs)

    def test_creates_output_directory(self, tmp_workspace: Path, simple_template: Path) -> None:
        """Create intermediate output directories if they do not exist."""
        from ci2lab.harness.tools.docx_writer import fill_docx_template

        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="subdir/nested/out.docx",
            fields={"name": "Test"},
        )

        assert "Error" not in result
        assert (tmp_workspace / "subdir" / "nested" / "out.docx").is_file()

    def test_success_message_contains_substitution_count(
        self, tmp_workspace: Path, simple_template: Path
    ) -> None:
        """The success message includes the number of substitutions."""
        from ci2lab.harness.tools.docx_writer import fill_docx_template

        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="count.docx",
            fields={"name": "X", "date": "Y", "city": "Z"},
        )

        assert "3 substitutions" in result or "3 substitution" in result

    def test_missing_template(self, tmp_workspace: Path) -> None:
        """Clear error if the template does not exist."""
        pytest.importorskip("docx", reason="python-docx not installed")
        from ci2lab.harness.tools.docx_writer import fill_docx_template

        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="does_not_exist.docx",
            output="out.docx",
            fields={},
        )

        assert result.startswith("Error")
        assert "does_not_exist.docx" in result

    def test_output_same_as_template_rejected(
        self, tmp_workspace: Path, simple_template: Path
    ) -> None:
        """The output cannot be the same path as the template."""
        from ci2lab.harness.tools.docx_writer import fill_docx_template

        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="templates/template.docx",
            fields={"name": "X"},
        )

        assert result.startswith("Error")
        assert "same path" in result

    def test_result_contains_kb(self, tmp_workspace: Path, simple_template: Path) -> None:
        """The success message includes the size in KB."""
        from ci2lab.harness.tools.docx_writer import fill_docx_template

        result = fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="kb_test.docx",
            fields={"name": "A"},
        )

        assert "KB" in result


# ---------------------------------------------------------------------------
# Tests de preview_fill_docx
# ---------------------------------------------------------------------------


class TestPreviewFillDocx:
    def test_valid_preview_shows_fields(self, tmp_workspace: Path, simple_template: Path) -> None:
        """The preview shows the template, the output and the fields."""
        from ci2lab.harness.tools.docx_writer import preview_fill_docx

        preview = preview_fill_docx(
            str(tmp_workspace),
            {
                "template": "templates/template.docx",
                "output": "output/result.docx",
                "fields": {"name": "Clara", "date": "2026"},
            },
        )

        assert preview.is_valid
        content = preview.new_content or ""
        assert "name" in content
        assert "Clara" in content
        assert "date" in content

    def test_missing_template_invalid(self, tmp_workspace: Path) -> None:
        """Invalid preview if the template does not exist."""
        pytest.importorskip("docx", reason="python-docx not installed")
        from ci2lab.harness.tools.docx_writer import preview_fill_docx

        preview = preview_fill_docx(
            str(tmp_workspace),
            {"template": "ghost.docx", "output": "out.docx", "fields": {}},
        )

        assert not preview.is_valid
        assert "ghost.docx" in (preview.validation_error or "")

    def test_no_template_param_invalid(self, tmp_workspace: Path) -> None:
        """Invalid preview if the template parameter is missing."""
        pytest.importorskip("docx", reason="python-docx not installed")
        from ci2lab.harness.tools.docx_writer import preview_fill_docx

        preview = preview_fill_docx(str(tmp_workspace), {"output": "out.docx", "fields": {}})

        assert not preview.is_valid

    def test_non_docx_template_invalid(self, tmp_workspace: Path) -> None:
        """Invalid preview if the template is not .docx."""
        pytest.importorskip("docx", reason="python-docx not installed")
        (tmp_workspace / "template.pdf").touch()
        from ci2lab.harness.tools.docx_writer import preview_fill_docx

        preview = preview_fill_docx(
            str(tmp_workspace),
            {"template": "template.pdf", "output": "out.docx", "fields": {}},
        )

        assert not preview.is_valid
        assert ".pdf" in (preview.validation_error or "")

    def test_new_file_flag(self, tmp_workspace: Path, simple_template: Path) -> None:
        """is_new_file=True if the output does not exist yet."""
        from ci2lab.harness.tools.docx_writer import preview_fill_docx

        preview = preview_fill_docx(
            str(tmp_workspace),
            {"template": "templates/template.docx", "output": "new.docx", "fields": {}},
        )

        assert preview.is_valid
        assert preview.is_new_file

    def test_existing_file_flag(self, tmp_workspace: Path, simple_template: Path) -> None:
        """is_new_file=False if the output already exists."""
        from ci2lab.harness.tools.docx_writer import fill_docx_template, preview_fill_docx

        fill_docx_template(
            cwd=str(tmp_workspace),
            template="templates/template.docx",
            output="existing.docx",
            fields={},
        )

        preview = preview_fill_docx(
            str(tmp_workspace),
            {"template": "templates/template.docx", "output": "existing.docx", "fields": {}},
        )

        assert preview.is_valid
        assert not preview.is_new_file


# ---------------------------------------------------------------------------
# Tests de _replace_in_paragraph
# ---------------------------------------------------------------------------


class TestReplaceInParagraph:
    def test_simple_replacement(self) -> None:
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document

        from ci2lab.harness.tools.docx_writer import _replace_in_paragraph

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello {{name}}, welcome to {{city}}.")

        count = _replace_in_paragraph(para, {"name": "Ana", "city": "Seville"})

        assert count == 2
        assert para.runs[0].text == "Hello Ana, welcome to Seville."

    def test_no_match_returns_zero(self) -> None:
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document

        from ci2lab.harness.tools.docx_writer import _replace_in_paragraph

        doc = Document()
        para = doc.add_paragraph("Text without placeholders.")

        count = _replace_in_paragraph(para, {"name": "Ana"})

        assert count == 0
        assert para.text == "Text without placeholders."

"""
fill_docx_template tool: generates a DOCX from a template
and a dictionary of {{placeholder}} → value fields.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from ci2lab.harness.tools.paths import PathViolationError, resolve_path
from ci2lab.harness.tools.write_preview import WritePreview

# ---------------------------------------------------------------------------
# Preview (without executing anything on disk)
# ---------------------------------------------------------------------------


def preview_fill_docx(cwd: str, args: dict[str, Any]) -> WritePreview:
    """Build a descriptive :class:`WritePreview` for ``fill_docx_template``.

    Does not read or write anything to disk; it only validates that the template
    exists and generates the confirmation text the user will see.

    Args:
        cwd: The current working directory used to resolve paths.
        args: Tool arguments; recognised keys are ``template`` (str),
            ``output`` (str) and ``fields`` (mapping of placeholder to value).

    Returns:
        A :class:`WritePreview` describing the planned write, carrying a
        ``validation_error`` when the template/output paths are invalid.
    """
    template = str(args.get("template", "")).strip()
    output = str(args.get("output", "")).strip()
    fields: dict[str, str] = {str(k): str(v) for k, v in (args.get("fields") or {}).items()}

    if not template:
        return WritePreview(
            path=output or "(no output)",
            is_new_file=True,
            diff="",
            validation_error="Error: 'template' parameter required",
        )
    if not output:
        return WritePreview(
            path="(no output)",
            is_new_file=True,
            diff="",
            validation_error="Error: 'output' parameter required",
        )

    # Validate that the template exists
    try:
        template_path = resolve_path(template, cwd)
    except PathViolationError as exc:
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error=str(exc),
        )

    if not template_path.is_file():
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error=f"Error: template not found: {template}",
        )
    if template_path.suffix.lower() != ".docx":
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error=(f"Error: the template must be a .docx, not '{template_path.suffix}'"),
        )

    # Validate that the output does not match the template
    try:
        output_path = resolve_path(output, cwd)
    except PathViolationError as exc:
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error=str(exc),
        )

    if output_path == template_path:
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error="Error: output cannot be the same path as the template",
        )

    if output_path.suffix.lower() != ".docx":
        return WritePreview(
            path=output,
            is_new_file=True,
            diff="",
            validation_error=(
                f"Error: the output must have a .docx extension, not '{output_path.suffix}'"
            ),
        )

    # Build descriptive text for the preview
    lines = [
        f"Template : {template}",
        f"Output    : {output} ({'existing — will be overwritten' if output_path.is_file() else 'new file'})",
        "",
    ]
    if fields:
        lines.append("Fields to substitute:")
        for key, value in fields.items():
            lines.append(f"  {{{{{key}}}}}  →  {value}")
    else:
        lines.append("(no fields provided — the template will be copied unchanged)")

    return WritePreview(
        path=str(output_path.relative_to(Path(cwd).resolve()))
        if output_path.is_relative_to(Path(cwd).resolve())
        else output,
        is_new_file=not output_path.is_file(),
        diff="",
        new_content="\n".join(lines),
    )


# ---------------------------------------------------------------------------
# Placeholder substitution in a DOCX paragraph
# ---------------------------------------------------------------------------


def _replace_in_paragraph(paragraph: Any, fields: dict[str, str]) -> int:
    """Substitute ``{{key}}`` → value in a paragraph, handling split runs.

    python-docx may split a placeholder's text across several runs
    (text fragments with different formatting). The strategy:
      1. Get the full text of the paragraph.
      2. Substitute in the full text.
      3. If there were changes, put the result in the first run and empty the rest.

    This preserves the formatting of the first run (font, bold, size)
    but loses formatting variations within the same paragraph. Acceptable in v1.

    Args:
        paragraph: A python-docx paragraph object to mutate in place.
        fields: Mapping of placeholder keys to replacement values.

    Returns:
        The number of distinct placeholder keys that were substituted.
    """
    full_text = paragraph.text
    new_text = full_text
    count = 0
    for key, value in fields.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))
            count += 1

    if count == 0:
        return 0

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)

    return count


# ---------------------------------------------------------------------------
# Main function
# ---------------------------------------------------------------------------


def fill_docx_template(
    cwd: str,
    template: str,
    output: str,
    fields: dict[str, str],
) -> str:
    """Generate a DOCX by filling a template with the given fields.

    Args:
        cwd: The current working directory used to resolve paths.
        template: Workspace-relative path to the ``.docx`` template.
        output: Workspace-relative destination path for the generated ``.docx``.
        fields: Mapping of placeholder keys to replacement values.

    Returns:
        A success message describing the generated document and substitution
        count, or an ``"Error: ..."`` message on any failure.
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: missing python-docx dependency. Run: pip install -e '.[convert]'"

    # Resolve paths (already validated in preview, but re-validated for safety)
    try:
        template_path = resolve_path(template, cwd)
        output_path = resolve_path(output, cwd)
    except PathViolationError as exc:
        return f"Error: {exc}"

    if not template_path.is_file():
        return f"Error: template not found: {template}"
    if template_path == output_path:
        return "Error: output cannot be the same path as the template"

    # Open the template
    try:
        doc = Document(str(template_path))
    except Exception as exc:
        return f"Error: could not open the template {template}: {exc}"

    # Create output directory if it does not exist
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Substitute placeholders in standalone paragraphs
    total = 0
    for paragraph in doc.paragraphs:
        total += _replace_in_paragraph(paragraph, fields)

    # Substitute placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += _replace_in_paragraph(paragraph, fields)

    # Save
    try:
        doc.save(str(output_path))
    except Exception as exc:
        return f"Error: could not save the document to {output}: {exc}"

    size_kb = round(output_path.stat().st_size / 1024, 1)
    return (
        f"Document generated: {output} "
        f"({total} substitution{'s' if total != 1 else ''}, {size_kb} KB)"
    )
